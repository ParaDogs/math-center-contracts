import math
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches
import datetime
import calendar
from calendar import Calendar
import PySimpleGUI as sg
import ru_number_to_text as ntt
import pymorphy2 as pm2
import pymorphy2_dicts_ru
import pickle
from pathlib import Path
educational_programms = {}
holydays_list = []

# при сборке (в pyinstaller) в исполняемый файл подключить C:/Python310/Lib/site-packages в --paths и C:/Python310/Lib/site-packages/pymorphy2_dicts_ru и contract_template.docx

# дописать отображение в создании программы

# загружает программы и выходные
def init():
    global educational_programms
    global holydays_list
    # подгрузить файл со всеми образовательными программами
    edu_dump_file = Path("educational_programms.dump")
    if edu_dump_file.is_file():
        with open(edu_dump_file.absolute(), 'rb') as handle:
            educational_programms = pickle.load(handle)
    else:
        with open(edu_dump_file.absolute(), 'wb') as handle:
            pickle.dump(educational_programms, handle, protocol=pickle.HIGHEST_PROTOCOL)
    # подгрузить файл с выходными
    hd_dump_file = Path("holydays_list.dump")
    if hd_dump_file.is_file():
        with open(hd_dump_file.absolute(), 'rb') as handle:
            holydays_list = pickle.load(handle)
    else:
        with open(hd_dump_file.absolute(), 'wb') as handle:
            pickle.dump(holydays_list, handle, protocol=pickle.HIGHEST_PROTOCOL)

# сохраняет программы и выходные
def save():
    global educational_programms
    global holydays_list
    # сохранить файл со всеми образовательными программами
    edu_dump_file = Path("educational_programms.dump")
    with open(edu_dump_file.absolute(), 'wb') as handle:
        pickle.dump(educational_programms, handle, protocol=pickle.HIGHEST_PROTOCOL)
    # сохранить файл с выходными
    hd_dump_file = Path("holydays_list.dump")
    with open(hd_dump_file.absolute(), 'wb') as handle:
        pickle.dump(holydays_list, handle, protocol=pickle.HIGHEST_PROTOCOL)

def gui():
    global educational_programms
    global holydays_list

    sg.theme('DarkBlue3')

    def update_calendar(year, month, days):
        window['Month'].update(f'{months_in_year[month-1]}')
        window['Year'].update(f'{year}')
        for row in range(6):
            for col in range(7):
                m, d = days[row * 7 + col]
                window[('Date', row, col)].update(d, text_color=white if m==month else gray,
                    background_color='blue' if d !='' and f"{str(int(d)).rjust(2,'0')}/{str(m).rjust(2,'0')}/{str(year).rjust(4,'0')}" in selected else sg.theme_background_color())
                window[('Date', row, col)].metadata=False if m==month else 'gray'

    months_in_year = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']

    today = datetime.datetime.now()
    year, month, day = today.year, today.month, today.day

    calendar = Calendar(firstweekday=0) # First week day is monday
    days = [(d.month, f'{d.day:>2d}') for d in calendar.itermonthdates(year=year, month=month)]
    days += [(None, '')] * (42-len(days))

    white, gray = '#FFFFFF', '#8699AF'
    bg = [sg.theme_background_color(), 'blue']

    selected = set()

    # вкладка создания договора
    main_layout = [
        [sg.Text('Образовательная программа')],
        [sg.Listbox([*educational_programms], size=(65,10), key='-COURSE-')],
        [sg.Text('Класс'), sg.InputText('1', size=(10,None), key='-GRADE-')],
        [sg.Text('Начало занятий'), sg.In(key='-STARTDATE-', enable_events=True, visible=True, readonly=True, size=(10,None)), sg.CalendarButton("Выбрать дату", close_when_date_chosen=False,begin_at_sunday_plus=1,format='%d/%m/%Y', target='-STARTDATE-')],
        [sg.Text('Конец занятий  '), sg.In(key='-ENDDATE-', enable_events=True, visible=True, readonly=True, size=(10,None)), sg.CalendarButton("Выбрать дату", close_when_date_chosen=False,begin_at_sunday_plus=1,format='%d/%m/%Y', target='-ENDDATE-',)],
        [sg.Text('Дни недели'), sg.Checkbox('ПН', key='-MON-'), sg.Checkbox('ВТ', key='-TUE-'), sg.Checkbox('СР', key='-WED-'), sg.Checkbox('ЧТ', key='-THU-'), sg.Checkbox('ПТ', key='-FRI-'), sg.Checkbox('СБ', key='-SAT-'), sg.Checkbox('ВС', key='-SUN-')],
        [sg.Text('Формат занятий'), sg.Radio('Индивидуально', 0, True,key='-INDIVIDUAL-'), sg.Radio('В группе', 0, False,key='-GROUP-')],
        [sg.Text('Место проведения'), sg.Radio('Очно', 1, True, key='-OFFLINE-'), sg.Radio('Онлайн', 1, False, key='-ONLINE-')],
        [sg.Text('Адрес проведения'), sg.Radio('1 речка', 2, False, key='-ADDRESS-'), sg.Radio('2 речка', 2, True, key='-ADDRESS1-')],
        

        [sg.Text('Данные ученика')],
        [sg.Text('ФИО:     '), sg.Input(key='-STUDENT-')],
        [sg.Text('Телефон:'), sg.Input(key='-STUDENTPHONE-')],

        [sg.Text('Данные родителя')],
        [sg.Text('ФИО:     '), sg.Input(key='-CUSTOMER-')],
        [sg.Text('Телефон:'), sg.Input(key='-CUSTOMERPHONE-')],

        [sg.Button('Сохранить договор в формате docx',pad=(120  ,0))]
    ]

    # вкладка создания новой программы
    edprogram_layout = [
        [sg.Text('Название образовательной программы')],
        [sg.InputText('', key='-COURSENAME-')],
        [sg.Text('Название программы в договоре')],
        [sg.InputText('', key='-CONTRACTCOURSENAME-')],
        [sg.Text('Время на одно занятие (в минутах)')],
        [sg.Radio('45', 3, True, key='-45min-'), sg.Radio('60', 3, False, key='-60min-'), sg.Radio('90', 3, False, key='-90min-'), sg.Radio('135', 3, False, key='-135min-')],
        # [sg.InputText('', key='-LESSONTIME-'), sg.Text('ак.ч.')],
        [sg.Text('Стоимость академического часа (45 мин.)')],
        [sg.InputText('', key='-HOURPRICE-'), sg.Text('руб.')],
        [sg.Text('Сообщения:')],
        [sg.Text(key='-EDPROGRAMOUTPUT-', size=(50,10), background_color='white', text_color='black')],
        [sg.Button('Сохранить новую образовательную программу',pad=(120, 0))],
        [sg.Button('Посмотреть все образовательные программы',pad=(120, 0))],
        [sg.Button('Очистить список образовательных программ',pad=(120, 0))]]

    # вкладка добавления выходных
    holydays_layout = [
        [sg.Text(f'{year}', key='Year'),
         sg.Text(sg.SYMBOL_DOWN, size=3, justification='center', background_color='green', enable_events=True, key='Year_DN'),
         sg.Text(sg.SYMBOL_UP, size=3, justification='center', background_color='green', enable_events=True, key='Year_UP'),
        #  sg.Push(),
         sg.Text(f'{months_in_year[month-1]}', key='Month'),
         sg.Text(sg.SYMBOL_DOWN, size=3, justification='center', background_color='green', enable_events=True, key='Month_DN'),
         sg.Text(sg.SYMBOL_UP, size=3, justification='center', background_color='green', enable_events=True, key='Month_UP')],
        [sg.Text()],
        [sg.Text(weekday, size=4, justification='center')
            for weekday in ('ПН', 'ВТ', 'СР', 'ЧТ', 'ПТ', 'СБ', 'ВС')]
    ]

    weeks = []
    for row in range(6):
        week = []
        for col in range(7):
            m, d = days[row * 7 + col]
            week.append(
                sg.Text(d, size=4, justification='center',
                    text_color=white if m==month else gray,
                    background_color='blue' if f"{str(int(d)).rjust(2,'0') if d != '' else '0'}/{str(m).rjust(2,'0')}/{str(year).rjust(4,'0')}" in selected else sg.theme_background_color(),
                    enable_events=True,
                    metadata=False if m==month else 'gray', key=('Date', row, col)))
        weeks.append(week)
    holydays_layout += weeks + [[sg.Button('Подтвердить выбор')]]
    holydays_layout += [[sg.Text('Сообщения:')],
        [sg.Text(key='-HOLYDAYSOUTPUT-', size=(50,10), background_color='white', text_color='black')],
        [sg.Button('Добавить выбранные дни к выходным', pad=(120,0))],
        [sg.Button('Посмотреть текущие выходные дни', pad=(120,0))],
        [sg.Button('Очистить список выходных дней', pad=(120,0))]]

    # общий слой, который представляет из себя группу вкладок
    layout = [
        [sg.TabGroup([[
            sg.Tab('Новый договор', main_layout),
            sg.Tab('Образовательные программы', edprogram_layout),
            sg.Tab('Праздники', holydays_layout)]],
            tab_location='top')]
    ]

    window = sg.Window('Договоры', layout)
    while True:
        event, values = window.read()
        if event in (None, 'Exit', 'Cancel', sg.WIN_CLOSED):
            break
        if event == 'Сохранить договор в формате docx':
            weekdays = dict(zip(('Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday'), tuple([values["-MON-"],
                values["-TUE-"],
                values["-WED-"],
                values["-THU-"],
                values["-FRI-"],
                values["-SAT-"],
                values["-SUN-"]])))
            if (sum(weekdays.values()) > 0 and
                values["-CUSTOMER-"] != "" and
                values["-STUDENT-"] != "" and
                values["-COURSE-"] != None and
                values["-GRADE-"] != None and
                values["-STARTDATE-"] != None and
                len(values["-COURSE-"]) != 0 and
                values["-ENDDATE-"] != None and
                is_name(values["-CUSTOMER-"]) and
                is_class_number(values["-GRADE-"]) and
                is_name(values["-STUDENT-"]) and
                is_phone_number(values["-STUDENTPHONE-"]) and
                is_phone_number(values["-CUSTOMERPHONE-"])):
                make_docx(
                    values["-CUSTOMER-"],
                    values["-CUSTOMERPHONE-"],
                    values["-STUDENT-"],
                    values["-STUDENTPHONE-"],
                    values["-COURSE-"][0],
                    values["-GRADE-"],
                    values["-STARTDATE-"],
                    values["-ENDDATE-"],
                    values["-OFFLINE-"],
                    values["-ADDRESS-"],
                    weekdays
                    )
                make_schedule(values["-COURSE-"][0], weekdays, values["-STARTDATE-"], values["-ENDDATE-"])
                sg.Popup('Договор сохранен')
            else:
                sg.Popup("Не все необходимые данные введены корректно", title="Ошибка ввода")
        if event == 'Сохранить новую образовательную программу':
            lesson_time = ("45","60","90","135")[(values["-45min-"],values["-60min-"],values["-90min-"],values["-135min-"]).index(True)]
            educational_programms[values["-COURSENAME-"]] = {
                'lesson_time': lesson_time,
                'contract_course_name': values["-CONTRACTCOURSENAME-"],
                'hour_price': values["-HOURPRICE-"]
                }
            save()
            window['-COURSE-'].update([*educational_programms])
            window['-EDPROGRAMOUTPUT-'].update('Добавлена новая программа')
        if event == 'Посмотреть все образовательные программы':
            window['-EDPROGRAMOUTPUT-'].update('Все добавленные образовательные программы:\n'+str(educational_programms))
        if event == 'Очистить список образовательных программ':
            educational_programms = {}
            save()
            window['-COURSE-'].update([*educational_programms])
            window['-EDPROGRAMOUTPUT-'].update('Список образовательных программ очищен')
        if event == 'Добавить выбранные дни к выходным':
            window['-HOLYDAYSOUTPUT-'].update('Список выходных дней обновлен')
            holydays_list += list(selected)
        if event == 'Посмотреть текущие выходные дни':
            window['-HOLYDAYSOUTPUT-'].update('Все добавленные выходные дни:\n'+str(holydays_list))
        if event == 'Очистить список выходных дней':
            holydays_list = []
            save()
            window['-HOLYDAYSOUTPUT-'].update('Список выходных дней очищен')
        elif event == 'Подтвердить выбор':
            window['-HOLYDAYSOUTPUT-'].update('Были выбраны следующие даты:\n'+str(selected))
        elif isinstance(event, tuple) and event[0] == 'Date':
            _, row, col = event
            m, d = days[row * 7 + col]
            if window[event].metadata != 'gray':
                window[event].metadata = not window[event].metadata
                window[event].update(background_color=bg[window[event].metadata])
                if window[event].metadata:
                    selected.add(f"{str(int(d)).rjust(2,'0')}/{str(m).rjust(2,'0')}/{str(year).rjust(4,'0')}")
                else:
                    selected.remove(f"{str(int(d)).rjust(2,'0')}/{str(m).rjust(2,'0')}/{str(year).rjust(4,'0')}")
        elif event in ('Month_UP', 'Month_DN', 'Year_UP', 'Year_DN'):
            delta = -1 if event.endswith('UP') else 1
            if event.startswith('Month'):
                m = month + delta
                year, month = (year-1, 12) if m < 1 else (year+1, 1) if m > 12 else (year, m)
            else:
                year += delta
            days = [(d.month, f'{d.day:>2d}') for d in calendar.itermonthdates(year=year, month=month)]
            days += [(None, '')] * (42-len(days))
            update_calendar(year, month, days)
    window.close()

def month_number_to_string(month_number):
    match int(month_number):
        case 1: return 'января'
        case 2: return 'февраля'
        case 3: return 'марта'
        case 4: return 'апреля'
        case 5: return 'мая'
        case 6: return 'июня'
        case 7: return 'июля'
        case 8: return 'августа'
        case 9: return 'сентября'
        case 10: return 'октября'
        case 11: return 'ноября'
        case 12: return 'декабря'

def month_number_to_string_default(month_number):
    match int(month_number):
        case 1: return 'январь'
        case 2: return 'февраль'
        case 3: return 'март'
        case 4: return 'апрель'
        case 5: return 'май'
        case 6: return 'июнь'
        case 7: return 'июль'
        case 8: return 'август'
        case 9: return 'сентябрь'
        case 10: return 'октябрь'
        case 11: return 'ноябрь'
        case 12: return 'декабрь'

def weekday_count(start, end):
  start_date = datetime.datetime.strptime(start, '%d/%m/%Y')
  end_date = datetime.datetime.strptime(end, '%d/%m/%Y')
  week = {}
  for i in range(-1,(end_date - start_date).days):
    day = calendar.day_name[(start_date + datetime.timedelta(days=i+1)).weekday()]
    week[day] = (week[day] + ((start_date + datetime.timedelta(days=i+1)).strftime('%d/%m/%Y') not in holydays_list)) if day in week else ((start_date + datetime.timedelta(days=i+1)).strftime('%d/%m/%Y') not in holydays_list)
  return week

def is_number(s):
    try:
        int(s)
        return True
    except:
        return False

def is_class_number(s):
    return is_number(s) and 1 <= int(s) <= 11

def is_phone_number(s):
    return math.prod([(sym in ('0','1','2','3','4','5','6','7','8','9','-','+',')','(',' ')) for sym in s]) != 0

def is_name(s):
    return math.prod(sym.isalpha() or sym in (' ','-') for sym in s) != 0

def set_col_widths(table):
    widths = (Cm(2.25), Cm(8.75), Cm(2.25), Cm(2.25))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def make_docx(customer, customer_phone, student, student_phone, course, grade, start, end, is_offline, address, weekdays):
    global educational_programms
    global holydays_list
    morph = pm2.MorphAnalyzer(pymorphy2_dicts_ru.get_path())
    doc = DocxTemplate("contract_template.docx")
    start_date_list = start.split('/')
    end_date_list = end.split('/')

    lessons_count = sum(v for k,v in weekday_count(start, end).items() if weekdays[k])
    lesson_time = int(educational_programms[course]['lesson_time'])
    price = round(float(educational_programms[course]['hour_price'])*(lesson_time/45), 2)
    contract_course_name = educational_programms[course]['contract_course_name']
    total_price = lessons_count * price
    start_date = datetime.datetime.strptime(start, '%d/%m/%Y')
    end_date = datetime.datetime.strptime(end, '%d/%m/%Y')

    customer_gent = ""
    for w in customer.split():
        try:
            customer_gent += morph.parse(w)[0].inflect({'gent'}).word.capitalize() + ' '
        except:
            customer_gent += w + ' '
    customer_gent = customer_gent[:-1]

    student_datv = ""
    for w in student.split():
        try:
            student_datv += morph.parse(w)[0].inflect({'datv'}).word.capitalize() + ' '
        except:
            student_datv += w + ' '
    student_datv = student_datv[:-1]

    context = {
        "CONTRACT_NUMBER" : 1, # randint(1,100),
        "CONTRACT_DAY" : str(datetime.datetime.now().day),
        "CONTRACT_MONTH" : month_number_to_string(datetime.datetime.now().month),
        "CONTRACT_YEAR" : str(datetime.datetime.now().year),

        # ФИО заказчика из именительного падежа переводится в родительный
        "CUSTOMER" : customer_gent,
        "CUSTOMER_PHONE" : customer_phone,
        # ФИО слушателя из именительного падежа переводится в дательный
        "STUDENT" : student_datv,
        "STUDENT_PHONE" : student_phone,
        "COURSE" : contract_course_name,
        "GRADE" : grade,

        "START_DAY" : start_date_list[0],
        "START_MONTH" : month_number_to_string(start_date_list[1]),
        "START_YEAR" : start_date_list[2],
        "END_DAY" : end_date_list[0],
        "END_MONTH" : month_number_to_string(end_date_list[1]),
        "END_YEAR" : end_date_list[2],

        "PLACE" : "очной" if is_offline else "online",
        "ADDRESS": "Океанский проспект, 87а." if address else "г. Владивосток, ул. Русская, 46б, 2 этаж.",

        "TOTAL_PRICE" : total_price,
        "TOTAL_PRICE_STRING" : ntt.num2text(total_price) + ' руб.',

        "LESSONS_PER_WEEK" : sum(weekdays.values()),
        "LESSON_TIME" : str(lesson_time*45),
    }
    doc.render(context)
    doc.save(f"{str(datetime.datetime.now()).replace(':','.')}-contract.docx")

def make_schedule(course, weekdays, start, end):
    start_date = datetime.datetime.strptime(start, '%d/%m/%Y')
    end_date = datetime.datetime.strptime(end, '%d/%m/%Y')
    dates = dict()
    for i in range(-1,(end_date - start_date).days):
        date = start_date + datetime.timedelta(days=i+1)
        day = calendar.day_name[date.weekday()]
        if weekdays[day] and date.strftime('%d/%m/%Y') not in holydays_list:
            # dates[f"{str(date.month).rjust(2,'0')}/{date.year}"] = (dates[f"{str(date.month).rjust(2,'0')}/{date.year}"] + [date.strftime('%d/%m/%Y')]) if f"{str(date.month).rjust(2,'0')}/{date.year}" in dates else [date.strftime('%d/%m/%Y')]
            dates[f"{str(date.month).rjust(2,'0')}/{date.year}"] = (dates[f"{str(date.month).rjust(2,'0')}/{date.year}"] + [str(date.day)]) if f"{str(date.month).rjust(2,'0')}/{date.year}" in dates else [str(date.day)]
    
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # date = f"{datetime.datetime.now().day} {month_number_to_string(datetime.datetime.now().month)} {datetime.datetime.now().year}"
    # head = doc.add_paragraph(f'Приложение 1.\nк договору №____\nот {date}')
    head = doc.add_paragraph('Приложение 1.\nк договору №____\nот _________')
    # выравнивание посередине
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    head.bold = True

    doc.add_paragraph("Оплату следующего месяца необходимо производить до 27-го числа текущего месяца. Например, если Вам необходимо оплатить услуги за октябрь, то плату следует внести до 27 сентября.")

    months = len(dates)
    table = doc.add_table(rows=months+1, cols=4, style='Table Grid')
    table.autofit = False

    header = table.rows[0]
    header.cells[0].text = 'Месяц Год'
    header.cells[1].text = 'Дни занятий'
    header.cells[2].text = 'Кол-во занятий'
    header.cells[3].text = 'Сумма, руб.'

    i = 0
    lesson_time = int(educational_programms[course]['lesson_time'])
    price = round(float(educational_programms[course]['hour_price'])*(lesson_time/45), 2)
    for row in table.rows[1:]:
        s = str(list(dates.keys())[i]).split('/')
        row.cells[0].text = month_number_to_string_default(s[0]) + ' ' + s[1]
        row.cells[1].text = str(', '.join(map(str, dates[list(dates.keys())[i]])))
        row.cells[2].text = str(len(dates[list(dates.keys())[i]]))
        row.cells[3].text = str(len(dates[list(dates.keys())[i]])*price)
        i += 1

    set_col_widths(table)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\nИсполнитель Первухин М.А. ______\tЗаказчик __________ ______")
    doc.save(f"{str(datetime.datetime.now()).replace(':','.')}-schedule.docx")

def main():
    init()
    gui()
    save()

if __name__ == "__main__":
    main()